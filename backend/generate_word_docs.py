"""Generates a consolidated Word (.docx) document from the Markdown docs in
../docs, per the specification's 'Word Documentation' output artifact
(Sections 3 and 8). Uses python-docx as the default documentation tool.

Run with: python generate_word_docs.py
Output: ../docs/RetailIQ_Platform_Documentation.docx
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt

from app.core.logging_config import get_logger

logger = get_logger(__name__)

DOCS_DIR = Path(__file__).resolve().parents[1] / "docs"
OUTPUT_PATH = DOCS_DIR / "RetailIQ_Platform_Documentation.docx"

DOC_ORDER = [
    "BRD.md",
    "FRD.md",
    "SRS.md",
    "Architecture.md",
    "API_Guide.md",
    "Installation_Guide.md",
    "Deployment_Guide.md",
    "User_Manual.md",
    "Developer_Guide.md",
    "Release_Notes.md",
]


def _add_markdown_paragraph(document: Document, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        return

    heading_match = re.match(r"^(#{1,4})\s+(.*)", stripped)
    if heading_match:
        level = len(heading_match.group(1))
        document.add_heading(heading_match.group(2), level=min(level, 4))
        return

    if stripped.startswith(("- ", "* ")):
        document.add_paragraph(stripped[2:], style="List Bullet")
        return

    if re.match(r"^\d+\.\s", stripped):
        document.add_paragraph(re.sub(r"^\d+\.\s", "", stripped), style="List Number")
        return

    if stripped.startswith("```"):
        return  # skip code fence markers; content lines are added as plain text

    document.add_paragraph(stripped)


def generate() -> Path:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    document.add_heading("RetailIQ Platform - Consolidated Documentation", level=0)
    document.add_paragraph(
        "Auto-generated from the project's Markdown documentation set. "
        "See individual .md files in the docs/ folder for the source of truth."
    )

    for filename in DOC_ORDER:
        path = DOCS_DIR / filename
        if not path.exists():
            logger.warning("Skipping missing doc: %s", filename)
            continue

        document.add_page_break()
        for line in path.read_text(encoding="utf-8").splitlines():
            _add_markdown_paragraph(document, line)

    document.save(str(OUTPUT_PATH))
    logger.info("Word documentation generated at %s", OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    generate()
